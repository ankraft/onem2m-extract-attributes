#
#	extractAttributes.py
#
#	(c) 2021 by Andreas Kraft
#	License: BSD 3-Clause License. See the LICENSE file for further details.
#
#	Extract attribute shortnames and other information from oneM2M's specification Word documents.
#

from __future__ import annotations
import argparse, json, csv, os
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Set, Union
from docx import Document
from docx.table import Table
import docx.opc.exceptions
from unidecode import unidecode
from rich.console import Console
from rich.progress import Progress, TextColumn, BarColumn
from rich.console import Console
from rich.table import Table


# TODO Table/output sorting choices

@dataclass
class AttributeTable:
	headers:list
	attribute:int
	shortname:int
	occursIn:int
	filename:str
	category:str

@dataclass
class Attribute:
	"""	Datastruture for an attribute, including shortname, longname, category etc.
	"""
	shortname: str
	attribute: str
	occurences:int
	occursIn:Set
	categories:Set
	documents:Set

	def asDict(self) -> dict:
		"""	Return this dataclass as a dictionary.
		"""
		return 	{	'shortname'	:	self.shortname,
					'attribute'	:	self.attribute,
					'occursIn'	:	sorted([ v for v in self.occursIn ]),
					'categories':	sorted([ v for v in self.categories ]),
					'documents'	:	sorted([ v for v in self.documents ])
				}

Attributes = Dict[str, Attribute]

#	Rich console for pretty printing
console = Console()

#	List of AttributeTable that define the various table headers to find the shortname tables inside the documents, and the
# 	offsets for the shortname, attribute, etc columns. 
#
# 	The following definitions may need to be updated and extended when new tables are added to the specification documents.

attributeTables:list[AttributeTable] = [

	# TS-0004
	AttributeTable(headers=['Parameter Name', 'XSD long name', 'Occurs in', 'Short Name'],	attribute=1, shortname=3, occursIn=2,  filename='ts-0004', category='Primitive Parameters'),
	AttributeTable(headers=['Root Element Name', 'Occurs in', 'Short Name'], 				attribute=0, shortname=2, occursIn=1,  filename='ts-0004', category='Primitive Root Elements'),
	AttributeTable(headers=['Attribute Name', 'Occurs in', 'Short Name'], 					attribute=0, shortname=2, occursIn=1,  filename='ts-0004', category='Resource Attributes'),
	AttributeTable(headers=['Resource Type Name', 'Short Name'], 							attribute=0, shortname=1, occursIn=-1, filename='ts-0004', category='Resource Types'),
	AttributeTable(headers=['Member Name', 'Occurs in', 'Short Name'],						attribute=0, shortname=2, occursIn=1,  filename='ts-0004', category='Complex Data Types'),
	AttributeTable(headers=['Member Name', 'Short Name'],									attribute=0, shortname=1, occursIn=-1, filename='ts-0004', category='Trigger Payload Fields'),

	# TS-0022
	AttributeTable(headers=['Attribute Name', 'Occurs in', 'Short Name', 'Notes'], 			attribute=0, shortname=2, occursIn=1,  filename='ts-0022', category='Common and Field Device Configuration'),
	AttributeTable(headers=['Member Name', 'Occurs in', 'Short Name', 'Notes'],				attribute=0, shortname=2, occursIn=1,  filename='ts-0022', category='Complex Data Types'),
	AttributeTable(headers=['ResourceType Name', 'Short Name'], 							attribute=0, shortname=1, occursIn=-1, filename='ts-0022', category='Resource Types'),		# Circumventing a typo in TS-0022 

	# TS-0023
	AttributeTable(headers=['Resource Type Name', 'Short Name'], 							attribute=0, shortname=1, occursIn=-1, filename='ts-0023', category='Specialization type short names'),
	AttributeTable(headers=['Attribute Name', 'Occurs in', 'Short Name'], 					attribute=0, shortname=2, occursIn=1,  filename='ts-0004', category='Resource attribute short names'),
	AttributeTable(headers=['Argument Name', 'Occurs in', 'Short Name'],					attribute=0, shortname=2, occursIn=1,  filename='ts-0023', category='Resource attribute short names'),

	# TS-0032
	AttributeTable(headers=['Attribute Name', 'Short Name'], 								attribute=0, shortname=1, occursIn=-1, filename='ts-0032', category='Security-specific Resource Type Short Names'),
	AttributeTable(headers=['Attribute Name', 'Occurs in', 'Short Name', 'Notes'], 			attribute=0, shortname=2, occursIn=1,  filename='ts-0032', category='Security-specific oneM2M Attribute Short Names'),
	AttributeTable(headers=['Member Name', 'Occurs in', 'Short Name', 'Notes'], 			attribute=0, shortname=2, occursIn=1,  filename='ts-0032', category='Security-specific oneM2M Complex data type member short names'),

]



def findAttributeTable(table:Table, filename:str) -> Union[AttributeTable, None]:
	"""	Search and return a fitting AttributeTable for the given document table.
		Return `None` if no fitting entry could be found.
	"""
	try:
		fn = filename.lower()
		row0 = table.rows[0]
		for snt in attributeTables:
			if len(snt.headers) != len(row0.cells):
				continue
			idx = 0
			isMatch = True
			while isMatch and idx < len(snt.headers):
				isMatch = row0.cells[idx].text == snt.headers[idx]
				idx += 1
			isMatch = isMatch and fn.startswith(snt.filename)
			if isMatch:
				return snt
	except:
		pass
	return None


def processDocuments(documents:list[str], outDirectory:str, csvOut:bool) -> Attributes|None:

	docs 							= {}
	ptasks 							= {}
	snCount							= 0
	attributes:dict[str, Attribute]	= {}

	with Progress(	TextColumn('[progress.description]{task.description}'),
					BarColumn(),
					TextColumn('[progress.percentage]{task.percentage:>3.0f}%'),
					speed_estimate_period=2.0) as progress:
		
		def stopProgress(msg:str='') -> None:
			progress.stop()
			progress.remove_task(readTask)
			console.print(msg)


		# Preparing tasks for progress
		readTask 	= progress.add_task(f'Reading document{"s" if len(documents)>1 else ""} ...', total=len(documents))

		#
		#	Read the input documents and add tasks for each of them
		#

		for d in documents:
			if not (dp := Path(d)).exists():
				stopProgress(f'[red]Input document "{d}" does not esist')
				return None
			if not dp.is_file():
				stopProgress(f'[red]Input document "{d}" is not a file')
				return None
			try:
				docs[d] = Document(d)
				ptasks[d] = progress.add_task(f'Processing {d} ...', total=1000)
				progress.update(readTask, advance=1)
			except docx.opc.exceptions.PackageNotFoundError as e:
				stopProgress(f'[red]Input document "{d}" is not a .docx file')
				return None
			except Exception as e:
				stopProgress(f'[red]Error reading file "{d}"')
				console.print_exception()
				return None
		
		# Add additional task
		checkTask	= progress.add_task('Checking results ...', total=2)
		writeTask	= progress.add_task('Writing files ...', total=2+len(documents) if csvOut else 2)

		#
		#	Process documents
		#

		for docName, doc in docs.items():
			processTask = ptasks[docName]

			# Process the document
			progress.update(processTask, total=len(doc.tables))
			for table in doc.tables:
				progress.update(processTask, advance=1)
				if (snt := findAttributeTable(table, docName)) is None:
					continue
				headersLen = len(snt.headers)
				for r in table.rows[1:]:
					cells = r.cells
					if cells[0].text.lower().startswith('note:') or len(r.cells) != headersLen:	# Skip note cells
						continue

					# Extract names and do a bit of transformations
					attributeName	= unidecode(cells[snt.attribute].text).strip()
					shortname 		= unidecode(cells[snt.shortname].text.replace('*', '').lower())
					occursIn 		= map(str.strip, unidecode(cells[snt.occursIn].text).split(',')) if snt.occursIn > -1 else ['n/a']	# Split and strip 'occurs in' entries
					
					# Don't process empty shortnames
					if not shortname:	
						continue

					# Create or update entry for shortname
					if shortname in attributes:
						entry = attributes[shortname]
						for v in occursIn:
							entry.occursIn.add(v)
						entry.categories.add(snt.category)
						entry.documents.add(docName)
						entry.occurences += 1
					else:
						entry = Attribute(	shortname=shortname,
											attribute=attributeName,
											occurences=1,
											occursIn=set([ v for v in occursIn ]),
											categories=set([ snt.category ]),
											documents=set([ docName ])
										)
					
					attributes[shortname] = entry
					snCount += 1
				continue

		#
		#	Further tests
		#
		progress.update(checkTask, advance=1)

		# count duplicates
		countDuplicates = 0
		for shortname, attribute in attributes.items():
			countDuplicates += 1 if attribute.occurences > 1 else 0
		progress.update(checkTask, advance=1)

		#
		#	generate outputs
		#

		# Write JSON output to a file
		progress.update(writeTask, advance=1)
		with open(f'{outDirectory}{os.sep}attributes.json', 'w') as jsonFile:
			json.dump([ v.asDict() for v in attributes.values()], jsonFile, indent=4)

		# Write output to CSV files
		if csvOut:
			for docName, doc in docs.items():			# Individually for each input file
				progress.update(writeTask, advance=1)
				# write a sorted list of attribute / shortnames to a csv file
				with open(f'{outDirectory}{os.sep}{docName.rsplit(".", 1)[0] + ".csv"}', 'w') as csvFile:
					writer = csv.writer(csvFile)
					writer.writerow(['Attribute', 'Short Name'])
					writer.writerows(	
						sorted(
							[ [attr.attribute, attr.shortname] for attr in attributes.values() if docName in attr.documents ],
							key=lambda x: x[0].lower() ))	# type: ignore [index]

		progress.update(writeTask, advance=1)

		#
		# finished. print further infos
		#

		progress.stop()
		console.print(f'Processed short names:      {snCount}')
		if countDuplicates > 0:
			console.print(f'Duplicate definitions:      {countDuplicates}')

	return attributes


def printAttributeTables(attributes:Attributes, duplicatesOnly:bool=True) -> None:
	"""	Print the found attributes to the console. Optionally print only duplicate entries.
	"""
	table = Table(show_lines=True, border_style='grey27')
	table.add_column('attribute', no_wrap=True)
	table.add_column('shortname', no_wrap=True, min_width=6)
	table.add_column('category', no_wrap=False)
	table.add_column('document(s)', no_wrap=False)
	for sn in sorted((attributes.keys())):
		attribute = attributes[sn]
		if attribute.occurences > 1:
			table.add_row(attribute.attribute, sn, ', '.join(attribute.categories), f'[red]{", ".join(attribute.documents)}')
		elif not duplicatesOnly:
			table.add_row(attribute.attribute, sn, ', '.join(attribute.categories), ', '.join(attribute.documents))
	console.print(table)


def printAttributeCsv(attributes:Attributes, duplicatesOnly:bool=True, outDirectory:str=None) -> None:
	"""	Print the found attributes to a CSV file. Optionally print only duplicate entries.
	"""
	# Write attributes also to a csv file
	with open(f'{outDirectory}{os.sep}{"attributes" if not duplicatesOnly else "duplicates"}.csv', 'w') as csvFile:
		writer = csv.writer(csvFile)
		writer.writerow(['Attribute', 'Short Name', 'Categories', 'Documents'])
		for sn in sorted((attributes.keys())):
			attribute = attributes[sn]
			if attribute.occurences > 1:
				writer.writerow([attribute.attribute, sn, ','.join(attribute.categories), ','.join(attribute.documents)])
			elif not duplicatesOnly:
				writer.writerow([attribute.attribute, sn, ','.join(attribute.categories), ','.join(attribute.documents)])



if __name__ == '__main__':

	# Parse command line arguments
	parser = argparse.ArgumentParser(formatter_class=argparse.ArgumentDefaultsHelpFormatter)
	parser.add_argument('--outdir', '-o', action='store', dest='outDirectory', default='out', metavar='<output directory>',  help='specify output directory')
	parser.add_argument('--csv', '-c', action='store_true', dest='csvOut', default=False, help='additionally generate shortname csv files')
	
	listArgs = parser.add_mutually_exclusive_group()
	listArgs.add_argument('--list', '-l', action='store_true', dest='list', default=False, help='list all found attributes')
	listArgs.add_argument('--list-duplicates', '-ld', action='store_true', dest='listDuplicates', default=False, help='list only duplicate attributes')

	parser.add_argument('document', nargs='+', help='documents to parse')
	args = parser.parse_args()

	# Process documents and print output
	os.makedirs(args.outDirectory, exist_ok=True)
	if (attributes := processDocuments(sorted(args.document), args.outDirectory, args.csvOut)) is None:
		exit(1)
	if args.list or args.listDuplicates:
		printAttributeTables(attributes, args.listDuplicates)
		if args.csvOut:
			printAttributeCsv(attributes, args.listDuplicates, args.outDirectory)